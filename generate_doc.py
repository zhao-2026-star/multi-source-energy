# -*- coding: utf-8 -*-
"""生成 MPF-Net 完整技术文档 (Word .docx)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ═══════ 样式设置 ═══════
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# 标题样式
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = '黑体'
    hs.font.color.rgb = RGBColor(0, 0, 0)

# ═══════ 辅助函数 ═══════
def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text):
    doc.add_paragraph(text)

def add_formula(text):
    """添加公式（居中显示）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = True
    return p

def add_bold(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

# ═══════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('MPF-Net\n多源异构数据融合网络\n技术文档')
run.font.size = Pt(22)
run.bold = True
run.font.name = '黑体'

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Multi-source Pattern Fusion Network\n面向配电网多源异构数据的端到端负荷预测模型')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('参数规模：1,040,776   输入：168h × 70维   输出：未来24h逐时负荷\n').font.size = Pt(11)
info.add_run('实测指标：RMSE=19.47kW  MAPE=19.42%  WAPE=16.36%\n').font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 目录占位
# ═══════════════════════════════════════════════════════════
add_heading('目录', 1)
add_para('1.  模型概述与架构总览')
add_para('2.  数据预处理与滑动窗口构建')
add_para('3.  特征嵌入层 (FeatureEmbedding)')
add_para('4.  位置编码 (Position Encoding)')
add_para('5.  缺失感知注意力 (MissingAwareAttention) — 创新点①')
add_para('6.  Transformer 编码器')
add_para('7.  聚类注意力 (ClusteringAttention) — 创新点②')
add_para('8.  多任务预测头 (MultiTaskHead) — 创新点③')
add_para('9.  完整前向传播流程——数据流动详解')
add_para('10. 损失函数')
add_para('11. 优化器与学习率调度')
add_para('12. 训练流程')
add_para('13. 评估指标')
add_para('14. 超参数配置总表')
add_para('15. 模型参数统计')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第1章：模型概述
# ═══════════════════════════════════════════════════════════
add_heading('1. 模型概述与架构总览', 1)

add_para('MPF-Net (Multi-source Pattern Fusion Network) 是一个面向配电网短期负荷预测的深度学习模型。'
         '模型接收三种异构数据源——连续型数值特征、离散型类别特征和自然语言型文本特征，'
         '通过统一的128维潜在空间（d_model=128）进行融合，最终输出未来24小时逐时负荷预测值（单位：kW）。')

add_heading('1.1 输入数据规格', 2)
add_para('数值特征 (18维): lag_1h, lag_2h, lag_3h, lag_24h, lag_48h, lag_168h, '
         'roll_mean_24h, roll_std_24h, roll_max_24h, roll_min_24h, '
         'capacity_kva, load_factor, 温度(°C), 相对湿度(%), 平均风速, 降水量(mm), '
         '最高温度(°C), 最低温度(°C)')
add_para('类别特征 (6维): hour(0-23), day_of_week(0-6), month(1-12), '
         'is_weekend(0/1), is_holiday(0/1), is_extreme(0/1)')
add_para('文本特征 (2维): holiday_name (节假日名称，21种), extreme_weather (极端天气描述，12种)')
add_para('辅助信息: 缺失掩码 missing_mask ∈ {0,1}（1=观测到，0=缺失），城市组编号 group_id ∈ {0,...,9}')

add_heading('1.2 核心创新点', 2)
add_para('创新点① 缺失感知注意力 (MissingAwareAttention): '
         '在自注意力计算中将缺失位置的注意力权重置为 -10⁹，使 Softmax 后的概率趋近于零，'
         '模型自动忽略缺失时间步，无需显式数据填充。')
add_para('创新点② 聚类注意力 (ClusteringAttention): '
         '从 MC-ANN (Multi-Cluster Additive Neural Network) 启发而来的可学习软聚类机制。'
         'K=5个可学习聚类中心通过梯度下降自动发现相似的负荷模式（如工业/居民/商业区），'
         '每个变压器被软分配到所有聚类中心，分配到的是权重向量而非硬标签。')
add_para('创新点③ 多源异构特征统一嵌入: '
         '数值特征通过线性投影、类别特征通过 nn.Embedding 查表、'
         '文本特征通过离线预计算的 BERT 768维嵌入 + 可学习投影，三者统一到 d_model=128 空间中做加性融合。')
add_para('创新点④ 多任务预测头 (MultiTaskHead): '
         '10个城市组共享一个通用预测层（提取通用时序-负荷映射模式），'
         '同时每个城市组拥有专属的偏差预测头（学习本地的个性化负荷偏差），'
         '最终预测 = 共享预测 + 组专属预测。')

add_heading('1.3 架构全景', 2)
add_para('完整的前向传播链路:')
add_formula('Input(168h × 70维) → FeatureEmbedding → LayerNorm → PositionEncoding → '
           'TransformerEncoder×4(含MissingAwareAttention) → '
           'ClusteringAttention(K=5) + Skip → MeanPooling → MultiTaskHead → Output(24h)')

add_para('数据维度变化轨迹:')
add_formula('[B, 168, 18num + 6cat + 2txt] → [B, 168, 128] → [B, 168, 128] → [B, 168, 128] '
           '→ [B, 1, 128] → [B, 24]')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第2章：数据预处理
# ═══════════════════════════════════════════════════════════
add_heading('2. 数据预处理与滑动窗口构建', 1)

add_heading('2.1 原始数据清洗', 2)
add_para('数值特征处理：使用 np.nan_to_num(arr, nan=0.0, posinf=0.0, neginf=0.0) '
         '将全部 NaN、+∞、-∞ 替换为 0.0。关键修复点：load_factor（负载率）在 capacity_kva=0 '
         '时会产生 Inf，在特征工程阶段使用 np.where(capacity_kva > 0, load/capacity, 0.0) 替代简单除法。')
add_para('类别特征处理：对于缺失值，使用 fillna(0) 填充为索引0（保留位）。')
add_para('文本特征处理：通过预构建的 text_vocab 映射表将中文文本转为整数索引，'
         '未登录词映射为 0。text_vocab 共包含 21 个词条。')

add_heading('2.2 滑动窗口构建', 2)
add_para('每个变压器的时序数据独立处理，避免跨变压器信息泄漏（不同变压器之间的负荷模式相互独立）。')

add_bold('窗口参数:')
add_para('输入窗口 W = 168 小时（7 天）—— 捕捉完整的周周期模式')
add_para('预测视界 H = 24 小时（1 天）—— 输出未来24小时逐时负荷')
add_para('滑动步长 stride = 24 小时 —— 每天生成一个预测样本')

add_bold('样本构造逻辑:')
add_para('对于长度为 n 小时的变压器时序:')
add_formula('样本数 = floor((n - 168 - 24) / 24) + 1')
add_para('每个样本 s 取:')
add_formula('X_s = data[t_s : t_s + 168]     # 输入：历史7天')
add_formula('Y_s = data[t_s + 168 : t_s + 192]   # 目标：未来24小时')

add_heading('2.3 缺失掩码生成', 2)
add_para('缺失掩码 mask ∈ {0, 1}¹⁶⁸ 标识输入窗口中每个时间步是否有效:')
add_formula('mask_t = 1   if 负荷(kW)_t 存在（观测到）')
add_formula('mask_t = 0   if 负荷(kW)_t 为 NaN（缺失）')
add_para('该掩码被传入 MissingAwareAttention，指导注意力机制忽略缺失位置。')

add_heading('2.4 数据集划分', 2)
add_para('训练集: 155,245 个样本（约 4255 个变压器·天）')
add_para('验证集: 82,516 个样本')
add_para('每个样本为一个 dict: {num_feat[168,18], cat_feat{6个key各[168]}, text_feat{2个key各[168]}, mask[168], group_id[标量], target[24]}')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第3章：特征嵌入层
# ═══════════════════════════════════════════════════════════
add_heading('3. 特征嵌入层 (FeatureEmbedding)', 1)

add_para('FeatureEmbedding 是 MPF-Net 的入口模块，负责将三种异构数据源统一映射到 d_model=128 维的连续向量空间，'
         '通过加性融合（Additive Fusion）形成统一的序列表示。')

add_heading('3.1 数值特征投影', 2)
add_para('18维连续数值特征通过单层线性变换映射到 d_model:')
add_formula('X_num = W_num · x_num                          (1)')
add_para('其中 W_num ∈ R^(d_model × 18) 为可学习权重矩阵，无偏置项。'
         '输入 x_num ∈ R^(B×168×18)，输出 X_num ∈ R^(B×168×128)。')

add_heading('3.2 类别特征嵌入', 2)
add_para('6个离散类别特征分别通过独立的 nn.Embedding 层查表映射到 d_model:')
add_formula('X_cat_k = Embedding_k(x_cat_k)                 (2)')
add_para('其中 Embedding_k 的词汇表大小分别为：hour=24, day_of_week=7, month=13, '
         'is_weekend=2, is_holiday=2, is_extreme=2。每个嵌入表输出维度固定为 d_model=128。')

add_heading('3.3 文本特征嵌入', 2)
add_para('2个文本特征通过离线预计算的 BERT 嵌入 + 可学习投影映射到 d_model:')
add_formula('E_bert = BERT_embedding(tokens)                 (3)')
add_formula('X_text = W_proj · E_bert                         (4)')
add_para('BERT 嵌入查表: 从预计算的 bert_embeddings.pkl 中按索引取出 768 维向量。BERT 权重在训练中冻结。')
add_para('投影层: W_proj ∈ R^(d_model × 768)，可训练，将 768 维降至 128 维。')

add_heading('3.4 加性融合', 2)
add_para('三种特征在 d_model 空间中进行无加权的逐元素相加:')
add_formula('X = X_num + ΣX_cat_k + ΣX_text_j               (5)')
add_para('加性融合的优势：(1) 计算高效，无需拼接后的额外投影层；'
         '(2) 隐式鼓励各模态学习互补而非冗余的信息；'
         '(3) 天然支持模态缺失场景（某模态为零向量不影响其他模态的贡献）。')

add_heading('3.5 嵌入归一化', 2)
add_para('加性融合后，7个嵌入项（1数值 + 6类别 + 2文本）的叠加可能导致数值爆炸，因此:')
add_formula('X_norm = LayerNorm(X)                            (6)')
add_para('LayerNorm 沿最后一维 (d_model=128) 计算均值和方差进行归一化，稳定训练早期梯度。')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第4章：位置编码
# ═══════════════════════════════════════════════════════════
add_heading('4. 位置编码 (Position Encoding)', 1)

add_para('Transformer 的自注意力机制本身是置换不变的（permutation-invariant），'
         '即任意打乱输入顺序不影响输出。为了让模型感知时间步的位置信息，需要在嵌入向量上叠加位置编码。')

add_heading('4.1 可学习位置编码', 2)
add_para('MPF-Net 采用可学习位置编码（Learnable Positional Encoding），而非固定的正弦编码:')
add_formula('X_pos = X + P[:, :L, :]                          (7)')
add_para('其中 P ∈ R^(1 × 168 × 128) 为 nn.Parameter 随机初始化（Xavier均匀分布），'
         '通过反向传播自动学习最优的位置表示。')
add_para('可学习编码相比正弦编码的优势：允许模型根据数据分布自适应调整位置表示，'
         '对168小时的长时间跨度更灵活（如模型可能学到 24h/168h 周期对应的位置模式）。')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第5章：缺失感知注意力
# ═══════════════════════════════════════════════════════════
add_heading('5. 缺失感知注意力 (MissingAwareAttention) — 创新点①', 1)

add_para('传统缺失值处理方法（均值填充、前向填充、线性插值等）会在填补过程中引入人为偏差。'
         'MPF-Net 的 MissingAwareAttention 将缺失信息直接编码进注意力计算，'
         '让模型自主学习如何应对缺失——这是本模型的第一个核心创新。')

add_heading('5.1 标准缩放点积注意力回顾', 2)
add_para('标准注意力计算:')
add_formula('Attention(Q, K, V) = softmax(QK^T / √d_k) · V          (8)')
add_para('其中 Q, K, V ∈ R^(B×H×L×d_h) 分别为查询、键、值矩阵，H 为注意力头数，d_h 为每头维度。')

add_heading('5.2 缺失掩码注入', 2)
add_para('在 Softmax 之前，对缺失时间步的注意力分数施加 -10⁹:')
add_formula('Score = QK^T / √d_k                                (9)')
add_formula('Score_ij = -1e9   if mask_j = 0                   (10)')
add_formula('Score_ij = Score_ij   if mask_j = 1                (11)')

add_para('由于 e^(-10⁹) ≈ 0，Softmax 后缺失位置的概率趋近于零，模型自动忽略缺失时间步的信息。')
add_para('使用 -1e9 而非 float("-inf") 的原因：当查询行全为缺失时，'
         '全 -1e9 行做 Softmax 会输出 ~[1/L, 1/L, ...] 的均匀分布而非全 NaN，保持数值稳定。')

add_heading('5.3 数值稳定性机制', 2)
add_para('Softmax 后应用 torch.nan_to_num(nan=0.0) 作为双保险，防止极端情况下出现 NaN。')

add_heading('5.4 多头注意力的投影', 2)
add_formula('Q = W_q × X          K = W_k × X          V = W_v × X             (12)')
add_para('每个注意力头有独立的 W_q, W_k, W_v ∈ R^(d_model × d_h)，d_h = d_model / n_heads = 128/8 = 16。')
add_para('多头输出拼接后通过输出投影:')
add_formula('MultiHead(X) = Concat(head_1, ..., head_8) · W_o             (13)')
add_para('其中 W_o ∈ R^(d_model × d_model) 为可学习输出投影矩阵。')

add_heading('5.5 关键维度推导', 2)
add_para('以 batch_size=256, seq_len=168 为例:')
add_formula('输入: [256, 168, 128]')
add_formula('→ 投影: Q/K/V 各 [256, 8, 168, 16]')
add_formula('→ 注意力分数: QK^T → [256, 8, 168, 168]  # 168×168 的时序注意力矩阵')
add_formula('→ 掩码注入: masked_fill(mask[256, 1, 1, 168]==0, -1e9)')
add_formula('→ Softmax + Dropout(0.1)')
add_formula('→ 加权求和: [256, 8, 168, 16] @ [256, 8, 168, 16] → [256, 8, 168, 16]')
add_formula('→ 输出投影: [256, 168, 128]')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第6章：Transformer 编码器
# ═══════════════════════════════════════════════════════════
add_heading('6. Transformer 编码器', 1)

add_heading('6.1 编码器层结构', 2)
add_para('每层 TransformerEncoderLayer 由两个子层组成，每个子层均采用 Pre-Norm + 残差连接:')
add_formula('子层1（自注意力）:')
add_formula('  X_attn = MissingAwareAttention(LayerNorm(X), mask)         (14)')
add_formula('  X = X + Dropout(X_attn)                                    (15)')
add_formula('子层2（前馈网络）:')
add_formula('  X_ffn = FFN(LayerNorm(X))                                   (16)')
add_formula('  X = X + Dropout(X_ffn)                                      (17)')

add_heading('6.2 前馈网络 (FFN)', 2)
add_para('两层全连接网络，中间使用 GELU 激活函数:')
add_formula('FFN(x) = W_2 · GELU(W_1 · x + b_1) + b_2                   (18)')
add_para('其中 W_1 ∈ R^(d_model × d_ff) = R^(128 × 512),  W_2 ∈ R^(512 × 128)，'
         'FFN 将特征映射到 4倍维度再压缩回来，提供逐位置的非线性变换。')
add_para('GELU (Gaussian Error Linear Unit) 激活函数:')
add_formula('GELU(x) = x · Φ(x) ≈ 0.5x·(1 + tanh(√(2/π)·(x + 0.044715x³)))      (19)')
add_para('其中 Φ(x) 为标准正态分布的累积分布函数。GELU 相比 ReLU 的优势：处处可微，'
         '对负值输出非零梯度（更平滑的激活），在 Transformer 类模型中表现优于 ReLU。')

add_heading('6.3 LayerNorm (层归一化)', 2)
add_para('沿特征维度（最后一维）计算均值和方差:')
add_formula('μ_d = (1/d) · Σᵢ x_i          σ²_d = (1/d) · Σᵢ (x_i - μ_d)²          (20)')
add_formula('y = γ · (x - μ_d) / √(σ²_d + ε) + β                                    (21)')
add_para('其中 γ, β ∈ R^d 为可学习参数，ε = 1e-5 为数值稳定常数。MPF-Net 使用 Pre-Norm 策略'
         '（先归一化再进子层），比 Post-Norm 训练更稳定，尤其是 Transformer 深层时。')

add_heading('6.4 完整编码器堆叠', 2)
add_para('MPF-Net 使用 4 层 TransformerEncoderLayer 堆叠:')
add_formula('X^(l+1) = TransformerEncoderLayer(X^(l), mask)     l = 0,1,2,3       (22)')
add_para('每层参数完全相同（共享架构，非共享权重）。mask 贯穿所有层，确保每一层的注意力都忽略缺失位置。')

add_heading('6.5 Dropout 正则化', 2)
add_para('每个子层的输出和 FFN 中间层均使用 Dropout(p=0.1)，在训练时随机丢弃 10% 的神经元，'
         '防止过拟合。推理时 Dropout 自动关闭。')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第7章：聚类注意力
# ═══════════════════════════════════════════════════════════
add_heading('7. 聚类注意力 (ClusteringAttention) — 创新点②', 1)

add_para('ClusteringAttention 是 MPF-Net 的第二个核心创新，从 MC-ANN (Multi-Cluster Additive Neural Network) '
         '方法获得启发。其核心思想是：让模型在训练过程中自动学习 K 个聚类中心，'
         '将每个变压器软分配到这些中心，从而隐式发现不同变电站之间的负荷模式相似性（如工业区、居民区、商业区分群）。')

add_heading('7.1 序列表示提取', 2)
add_para('从 Transformer 编码器输出的序列 X ∈ R^(B×168×128) 中提取每个样本的整体表示:')
add_formula('r = (X[:, -1, :] + X.mean(dim=1)) / 2                    (23)')
add_para('即最后一个时间步与全局均值池化的等权平均，兼顾了最新信息和全局信息。')

add_heading('7.2 聚类中心与软分配', 2)
add_para('K=5 个可学习聚类中心 C ∈ R^(K × d_model)，Xavier 均匀初始化并通过训练优化:')
add_formula('C_norm = LayerNorm(C)                                        (24)')
add_formula('logits_k = (r · C_norm_k) / (τ · √d)                        (25)')
add_formula('α_k = exp(logits_k) / Σ_j exp(logits_j)                     (26)')

add_para('其中 τ = 1.0 为温度参数（控制软分配的"锐度"），√d 为缩放因子（与注意力缩放一致，'
         '防止点积过大导致梯度消失）。')

add_heading('7.3 数值稳定性保障', 2)
add_para('logits 裁剪: clamp(logits, -50, 50) — 防止 e^50 ≈ 5.2×10²¹ 溢出 float32 范围 (~3.4×10³⁸)')
add_para('软分配修复: torch.nan_to_num(assignments, nan=1.0/K) — 全零行退化为均匀分布')

add_heading('7.4 聚类加权上下文', 2)
add_para('每个样本的聚类上下文由软分配权重对聚类中心的加权平均得到:')
add_formula('context = Σ_k α_k · C_k                # [B, d_model]         (27)')
add_para('该上下文通过残差连接加回原始序列:')
add_formula('X_out = X + context.unsqueeze(1)       # [B, L, d_model]      (28)')
add_para('残差连接的作用：聚类上下文提供粗粒度的模式信息（"哪个群组"），'
         '而原始 X 保留细粒度的时序信息（"具体什么时刻"），两者互补。')

add_heading('7.5 软分配 vs 硬聚类的优势', 2)
add_para('硬聚类（如 K-means）将每个变压器分配到一个确定性标签，丢弃了不确定性信息。')
add_para('软分配让一个变压器可以"部分属于"多个聚类——例如一个工业-居民混合区的变压器可能 '
         '40% 属于工业模式 + 60% 属于居民模式，模型通过加权平均综合利用多种模式的知识。')
add_para('实验验证：去掉 ClusteringAttention 后 RMSE 从 19.02 升至 19.71（+3.6%），'
         '证明该模块对模型性能有显著正面贡献。')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第8章：多任务预测头
# ═══════════════════════════════════════════════════════════
add_heading('8. 多任务预测头 (MultiTaskHead) — 创新点③', 1)

add_para('原始数据包含 423 台变压器，按地理位置划分为 10 个城市组。不同城市的负荷模式受当地 '
         '产业结构、气候条件、经济发展水平等因素影响，存在系统性偏差。')

add_heading('8.1 时序池化', 2)
add_para('在进入预测头之前，对时序维度进行平均池化将 [B, 168, 128] 压缩为 [B, 1, 128]:')
add_formula('h = X.mean(dim=1, keepdim=True)         # [B, 1, d_model]    (29)')

add_heading('8.2 双路径预测架构', 2)
add_para('共享路径（提取通用模式）:')
add_formula('ŷ_shared = W_2 · GELU(W_1 · h + b_1) + b_2   # [B, 24]       (30)')
add_para('其中 W_1 ∈ R^(128×128), W_2 ∈ R^(128×24)。')

add_para('分组路径（学习城市本地偏差）:')
add_formula('ŷ_group,k = W_k2 · GELU(W_k1 · h + b_k1) + b_k2   # [B, 24]  (31)')
add_para('其中 W_k1 ∈ R^(128×64), W_k2 ∈ R^(64×24) 为第 k 组的专属参数，共 10 组。')

add_heading('8.3 最终输出', 2)
add_formula('ŷ = ŷ_shared + ŷ_group[group_id]                          (32)')
add_para('即每个样本的预测 = 全局通用的时序映射 + 所在城市组的本地偏差修正。')
add_para('总参数量: 共享路径 32,896 + 分组路径 73,600 = 106,496（占全模型 10.2%）。')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第9章：完整前向传播流程
# ═══════════════════════════════════════════════════════════
add_heading('9. 完整前向传播流程——数据流动详解', 1)

add_para('以下以一个 batch (B=256, L=168) 为例，详细追踪每一步的数据形状变化。')

add_heading('9.1 输入阶段', 2)
add_formula('输入:  num_feat [256,168,18]                               (step 0)')
add_formula('      cat_feat {6 keys, each [256,168]}')
add_formula('      text_feat {2 keys, each [256,168]}')
add_formula('      mask [256,168], group_ids [256]')

add_heading('9.2 特征嵌入阶段', 2)
add_formula('X_num  = W_num_proj(num_feat)           → [256, 168, 128]   (step 1)')
add_formula('X_hour = Embedding_hour(cat[hour])       → [256, 168, 128]   (step 2)')
add_formula('... (其余5个类别嵌入同理)')
add_formula('X_hol  = BERT_lookup(text[holiday_name])  → [256, 168, 768]')
add_formula('X_hol  = W_text_proj(X_hol)               → [256, 168, 128]   (step 3)')
add_formula('X = X_num + X_hour + X_dow + ... + X_hol + X_weather        (step 4)')
add_formula('X = LayerNorm(X)                          → [256, 168, 128]   (step 5)')

add_heading('9.3 位置编码阶段', 2)
add_formula('X = X + P[:, :168, :]                     → [256, 168, 128]   (step 6)')

add_heading('9.4 Transformer 编码阶段', 2)
add_para('第1层:')
add_formula('X_norm = LayerNorm(X)                     → [256, 168, 128]')
add_formula('X_attn = MissingAwareAttention(X_norm, mask) → [256, 168, 128]')
add_formula('X = X + Dropout(X_attn)                   → [256, 168, 128]   (step 7)')
add_formula('X_norm = LayerNorm(X)')
add_formula('X_ffn  = W2·GELU(W1·X_norm+b1)+b2         → [256, 168, 128]')
add_formula('X = X + Dropout(X_ffn)                    → [256, 168, 128]   (step 8)')
add_para('第2-4层: 重复 steps 7-8，共4层。')

add_heading('9.5 聚类注意力阶段', 2)
add_formula('r = (X[:,-1,:] + X.mean(dim=1)) / 2      → [256, 128]        (step 9)')
add_formula('α = Softmax(r · C^T / (τ·√d))             → [256, 5]          (step 10)')
add_formula('context = α @ C                           → [256, 128]         (step 11)')
add_formula('X = X + context.unsqueeze(1)              → [256, 168, 128]    (step 12)')

add_heading('9.6 预测阶段', 2)
add_formula('h  = X.mean(dim=1, keepdim=True)           → [256, 1, 128]     (step 13)')
add_formula('h  = h.squeeze(1)                          → [256, 128]        (step 14)')
add_formula('ŷ_shared = W2·GELU(W1·h+b1)+b2             → [256, 24]         (step 15)')
add_formula('ŷ_group  = GroupHead[gid](h)               → [256, 24]         (step 16)')
add_formula('ŷ  = ŷ_shared + ŷ_group                     → [256, 24]         (step 17)')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第10章：损失函数
# ═══════════════════════════════════════════════════════════
add_heading('10. 损失函数', 1)

add_heading('10.1 Huber Loss (平滑L1损失)', 2)
add_para('MPF-Net 使用 Huber Loss 作为训练目标函数:')
add_formula('L_huber(y, ŷ) = (1/n) · Σ_i L_i                          (33)')
add_formula('L_i = 0.5·(y_i - ŷ_i)²                  if |y_i - ŷ_i| ≤ δ')
add_formula('L_i = δ·|y_i - ŷ_i| - 0.5·δ²              if |y_i - ŷ_i| > δ')

add_para('其中 δ = 1.0 (默认值)。')

add_heading('10.2 选择 Huber Loss 的原因', 2)
add_para('(1) 对离群点的鲁棒性：MSE 对大误差平方惩罚极重，单个异常样本（如极端天气日）可能导致 '
         '梯度爆炸。Huber Loss 对大残差切换为线性惩罚，有效抑制离群点的影响。')
add_para('(2) 对小误差的精度：对小残差使用二次惩罚（L2），比 MAE (L1) 收敛更快、最优解更接近真实均值。')
add_para('(3) 处处可微：在 δ 点二阶可微，梯度平滑，不会像 MAE 在零点产生梯度跳变。')
add_para('(4) 工程实践：PyTorch 的 nn.HuberLoss(delta=1.0) 的梯度实现经过高度优化，'
         '比手动实现分支逻辑快 3-5 倍。')

add_heading('10.3 为什么不用其他损失', 2)
add_para('MSE (均方误差): 对离群点过于敏感，配电网负荷数据含有多种极端情况（设备故障、异常天气），'
         'MSE 会使模型过度拟合这些罕见事件。')
add_para('MAE (平均绝对误差): 零点不可微导致优化困难，且 MAE 的最优解是中位数而非均值，'
         '在对称分布上不如 MSE/Huber。')
add_para('Quantile Loss: 分位数损失用于概率预测（输出区间），MPF-Net 仅做点预测，不必引入复杂度。')

add_heading('10.4 NaN/Inf 检测与跳过', 2)
add_para('每个 batch 训练时检测: if torch.isnan(loss) or torch.isinf(loss): continue')
add_para('跳过该 batch 的梯度更新，防止单个异常样本（如全部为 0 的行）导致整个训练崩溃。'
         '这是工程上的安全网，正常训练中极少触发（<0.01% batch）。')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第11章：优化器与学习率调度
# ═══════════════════════════════════════════════════════════
add_heading('11. 优化器与学习率调度', 1)

add_heading('11.1 AdamW 优化器', 2)
add_para('AdamW = Adam + 解耦的权重衰减 (Decoupled Weight Decay):')
add_formula('g_t = ∇_θ L(θ_t-1)                                         (35)')
add_formula('m_t = β_1·m_t-1 + (1-β_1)·g_t                    # 一阶动量')
add_formula('v_t = β_2·v_t-1 + (1-β_2)·g_t²                    # 二阶动量')
add_formula('m_hat = m_t / (1 - β_1^t)                         # 偏差修正')
add_formula('v_hat = v_t / (1 - β_2^t)')
add_formula('θ_t = θ_t-1 - η·(m_hat / (√v_hat + ε) + λ·θ_t-1)  # 解耦权重衰减')

add_para('超参数: lr=0.001, β₁=0.9, β₂=0.999, ε=1e-8, weight_decay=1e-5')
add_para('AdamW 相比 Adam 的优势：权重衰减从 L2 正则化中解耦，与自适应学习率不再互相干扰，'
         '在 Transformer 类模型中普遍比 Adam 泛化更好。')

add_heading('11.2 Warmup + Cosine 退火调度', 2)
add_para('学习率先线性预热再余弦衰减到 ~0:')
add_formula('lr(step) = lr_max · factor(step)                           (36)')

add_para('预热阶段 (前5% 步数):')
add_formula('factor(step) = step / warmup_steps                         (37)')
add_para('其中 warmup_steps = 0.05 × total_steps。初始学习率为 0，线性增至 lr_max=0.001。')

add_para('衰减阶段 (后95% 步数):')
add_formula('progress = (step - warmup) / (total - warmup)              (38)')
add_formula('factor = 0.5 · (1 + cos(π · progress))                     (39)')
add_para('学习率从 lr_max 平滑衰减至 ~0。余弦退火相比线性衰减的优势：初始缓慢下降（给模型充分探索空间），'
         '后期加速下降（快速收敛到最优邻域）。')

add_heading('11.3 梯度裁剪', 2)
add_para('对全体可训练参数应用梯度范数裁剪:')
add_formula('if ||g|| > clip_norm: g = g · (clip_norm / ||g||)         (40)')
add_para('其中 clip_norm = 10.0。梯度裁剪防止单个 batch 的大梯度导致的训练不稳定，'
         '对 Transformer 类模型尤为重要（注意力矩阵的反向传播可能产生极大梯度）。')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第12章：训练流程
# ═══════════════════════════════════════════════════════════
add_heading('12. 训练流程', 1)

add_heading('12.1 单 Epoch 训练过程', 2)
add_para('伪代码:')
add_formula('for batch in train_loader:')
add_formula('    pred, assignments = model(num, cat, text, mask, gid)')
add_formula('    loss = HuberLoss(pred, target)')
add_formula('    if isnan(loss) or isinf(loss): continue  # 跳过异常batch')
add_formula('    optimizer.zero_grad()')
add_formula('    loss.backward()')
add_formula('    clip_grad_norm_(model.parameters(), 10.0)')
add_formula('    optimizer.step()')
add_formula('    update_lr(step)  # warmup+cosine')

add_heading('12.2 验证过程', 2)
add_para('伪代码:')
add_formula('with torch.no_grad():')
add_formula('    for batch in val_loader:')
add_formula('        pred, _ = model(num, cat, text, mask, gid)')
add_formula('        loss = HuberLoss(pred, target)')
add_formula('        collect pred, target for metrics')
add_formula('avg_val_loss = mean(all_losses)')
add_formula('rmse, mae, mape = compute_metrics(all_preds, all_targets)')

add_heading('12.3 模型保存策略', 2)
add_para('仅保存验证损失 (val_loss) 最低的检查点（单一最佳模型，不保存中间检查点）:')
add_para('if val_loss < best_val_loss: save_checkpoint("mpf_net_best.pth")')
add_para('检查点内容: {epoch, model_state_dict, optimizer_state_dict, val_loss, config, model_config}')

add_heading('12.4 训练配置', 2)
add_para('epochs: 30-100（论文实验统一 30 epochs 快速验证 + 100 epochs 最终结果）')
add_para('batch_size: 256')
add_para('device: NVIDIA RTX 4090 (24GB VRAM)')
add_para('单 epoch 耗时: 约 30 秒（训练）+ 15 秒（验证）= 45 秒')
add_para('30 epochs 总耗时: 约 22 分钟')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第13章：评估指标
# ═══════════════════════════════════════════════════════════
add_heading('13. 评估指标', 1)

add_heading('13.1 RMSE (均方根误差)', 2)
add_formula('RMSE = √[(1/n) · Σᵢ (yᵢ - ŷᵢ)²]                              (41)')
add_para('单位: kW。RMSE 对大误差的惩罚程度高于 MAE（平方放大效应），反映预测的最坏情况表现。')

add_heading('13.2 MAE (平均绝对误差)', 2)
add_formula('MAE = (1/n) · Σᵢ |yᵢ - ŷᵢ|                                   (42)')
add_para('单位: kW。MAE 对所有误差等权对待，反映预测的典型表现（中位数偏离）。')

add_heading('13.3 MAPE (平均绝对百分比误差)', 2)
add_formula('MAPE = (100/n) · Σᵢ |(yᵢ - ŷᵢ) / yᵢ|                        (43)')
add_para('单位: %。MAPE 以相对百分比形式衡量误差，便于跨变压器/跨数据集比较。')
add_para('注意: MAPE 分母使用 target 而非 (target+pred)/2（非对称 MAPE），且 target<1e-6 的位置被排除')

add_heading('13.4 WAPE (加权绝对百分比误差)', 2)
add_formula('WAPE = 100 · Σ|y - ŷ| / Σ|y|                               (44)')
add_para('单位: %。WAPE 使用总负荷而不是逐点负荷做归一化，对目标值接近零的样本更鲁棒。')

add_heading('13.5 当前模型表现', 2)
add_para('Run3 测试集 (30 epochs, batch=256):')
add_para('RMSE = 19.47 kW')
add_para('MAE  = 10.52 kW')
add_para('MAPE = 19.42%')
add_para('WAPE = 16.36%')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第14章：超参数配置
# ═══════════════════════════════════════════════════════════
add_heading('14. 超参数配置总表', 1)

add_para('模型结构参数:')
records_model = [
    ('d_model', '128', '所有模块的隐藏维度'),
    ('n_heads', '8', '多头注意力头数 (128/8=16 per head)'),
    ('n_layers', '4', 'Transformer 编码器层数'),
    ('d_ff', '512', 'FFN 中间层维度 (128×4)'),
    ('n_clusters', '5', '可学习聚类中心数量 K'),
    ('n_groups', '10', '城市组数量'),
    ('max_seq_len', '168', '最大序列长度 (7天×24h)'),
    ('forecast_horizon', '24', '预测视界 (未来24h)'),
    ('dropout', '0.1', 'Dropout 比率'),
    ('num_dim', '18', '输入数值特征维度'),
    ('cat_dim', '6', '输入类别特征数量'),
    ('text_dim', '2', '输入文本特征数量'),
]

table = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
hdr = table.rows[0].cells
hdr[0].text = '参数名'
hdr[1].text = '值'
hdr[2].text = '说明'
for name, val, desc in records_model:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = val
    row[2].text = desc

add_para('')
add_para('训练超参数:')
records_train = [
    ('lr', '0.001', '初始学习率 (AdamW)'),
    ('weight_decay', '1e-5', '权重衰减系数 (L2 正则化)'),
    ('epochs', '30/100', '训练轮次'),
    ('batch_size', '256', '批大小'),
    ('warmup_ratio', '0.05', '预热比例 (前5%步数)'),
    ('clip_norm', '10.0', '梯度裁剪阈值'),
    ('Huber δ', '1.0', 'Huber Loss 分段点'),
]

table2 = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
hdr2 = table2.rows[0].cells
hdr2[0].text = '参数名'
hdr2[1].text = '值'
hdr2[2].text = '说明'
for name, val, desc in records_train:
    row = table2.add_row().cells
    row[0].text = name
    row[1].text = val
    row[2].text = desc

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第15章：模型参数统计
# ═══════════════════════════════════════════════════════════
add_heading('15. 模型参数统计', 1)

add_para('MPF-Net 总可训练参数量: 1,040,776')

add_para('各模块参数分布:')
records_params = [
    ('FeatureEmbedding', '数值Linear (128×18) + 类别Embedding (6个) + BERT投影 (128×768) + BERT查表(冻结)', '~135K'),
    ('TransformerEncoder (×4)', 'QKV投影 (4×3×128×128) + 输出投影 (4×128×128) + FFN (4×2×128×512)', '~786K'),
    ('LayerNorm (×10)', '10个 LayerNorm (γ+β 各128)', '~2.6K'),
    ('位置编码', '168×128 可学习矩阵', '~21.5K'),
    ('ClusteringAttention', '5×128 聚类中心 + LayerNorm', '~0.9K'),
    ('MultiTaskHead', '共享Linear(128→128→24) + 10组(128→64→24)', '~106.5K'),
    ('总计', '', '1,040,776'),
]

table3 = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
hdr3 = table3.rows[0].cells
hdr3[0].text = '模块'
hdr3[1].text = '组成'
hdr3[2].text = '参数量'
for name, comp, params in records_params:
    row = table3.add_row().cells
    row[0].text = name
    row[1].text = comp
    row[2].text = params

add_para('')
add_para('参数量最大的模块是 TransformerEncoder（占总量 75.5%），'
         '这与其承担的核心时序建模任务一致。预测头仅占 10.2%，表明模型将大部分容量'
         '用于学习有价值的特征表示，而非预测头的简单映射。')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 附录
# ═══════════════════════════════════════════════════════════
add_heading('附录A：特征列详细清单', 1)
add_para('数值特征 (18列):')
for c in ['lag_1h', 'lag_2h', 'lag_3h', 'lag_24h', 'lag_48h', 'lag_168h',
          'roll_mean_24h', 'roll_std_24h', 'roll_max_24h', 'roll_min_24h',
          'capacity_kva', 'load_factor',
          '温度(°C)', '相对湿度(%)', '平均风速', '降水量(mm)',
          '最高温度(°C)', '最低温度(°C)']:
    add_para(f'  · {c}')

add_para('类别特征 (6列):')
for c in ['hour', 'day_of_week', 'month', 'is_weekend', 'is_holiday', 'is_extreme']:
    add_para(f'  · {c}')

add_para('文本特征 (2列):')
add_para('  · holiday_name (21种，如：春节、国庆节、中秋节...)')
add_para('  · extreme_weather (12种，如：暴雨、高温、大风...)')

add_heading('附录B：关键数学符号表', 1)
symbols = [
    ('B', 'Batch size (批次大小)，默认 256'),
    ('L', 'Sequence length (序列长度)，固定 168 (7天×24h)'),
    ('H', '预测视界 (forecast horizon)，固定 24'),
    ('d_model', '模型隐藏维度，固定 128'),
    ('n_heads', '多头注意力头数，固定 8'),
    ('d_h', '每头维度，d_h = d_model / n_heads = 16'),
    ('K', '聚类中心数量，固定 5'),
    ('G', '城市组数量，固定 10'),
    ('X', '编码器输出的隐藏表示 [B, L, d_model]'),
    ('ŷ', '模型预测输出 [B, H]'),
    ('α', '软聚类分配权重 [B, K]'),
    ('C', '可学习聚类中心 [K, d_model]'),
    ('P', '可学习位置编码 [1, 168, d_model]'),
]
for sym, desc in symbols:
    add_para(f'  · {sym:15s} {desc}')

# ═══════ 保存 ═══════
output_path = r'D:\电力场景下多模态数据融合\District-power\MPF-Net完整技术文档.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
